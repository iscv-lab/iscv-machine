from datetime import datetime
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import Pt
import os, sys
from docx2pdf import convert
import numpy as np
import matplotlib
import matplotlib.pyplot as plt
from matplotlib.patches import Circle, RegularPolygon
from matplotlib.path import Path
from matplotlib.projections.polar import PolarAxes
from matplotlib.projections import register_projection
from matplotlib.spines import Spine
from matplotlib.transforms import Affine2D
import json
import asyncio
import pandas as pd
from utils.string import to_str
from sklearn.preprocessing import MinMaxScaler
import subprocess
sys.path.append(os.path.abspath(os.path.join("..", "..")))
matplotlib.use("agg")


def ScoreMinMaxScaler(result, type):
    score = [
        [0],
        [1],
        [2],
        [3],
        [4],
        [5],
        [6],
        [7],
        [8],
        [9],
        [10],
        [11],
        [12],
        [13],
        [14],
        [15],
        [16],
        [17],
        [18],
        [19],
        [20],
        [21],
        [22],
        [23],
        [24],
        [25],
        [26],
        [27],
        [28],
        [29],
        [30],
        [31],
        [32],
        [33],
        [34],
        [35],
        [36],
        [37],
        [38],
        [39],
        [40],
    ]

    # Openness Scaler
    if type == "Openness":
        # O_scaler=MinMaxScaler(feature_range=(-3.2, 2.3))
        O_score_data = O_scaler.fit_transform(score)
        O_score_scale = O_score_data[result]
        return O_score_scale

    # Conscientiousness Scaler
    if type == "Conscientiousness":
        # C_scaler=MinMaxScaler(feature_range=(-2.5, 2.4))
        C_score_data = C_scaler.fit_transform(score)
        C_score_scale = C_score_data[result]
        return C_score_scale

    # Extroversion Scaler
    if type == "Extroversion":
        # E_scaler=MinMaxScaler(feature_range=(-3.6, 2.5))
        E_score_data = E_scaler.fit_transform(score)
        E_score_scale = E_score_data[result]
        return E_score_scale

    # Agreeableness Scaler
    if type == "Agreeableness":
        # A_scaler=MinMaxScaler(feature_range=(-3.0, 2.1))
        A_score_data = A_scaler.fit_transform(score)
        A_score_scale = A_score_data[result]
        return A_score_scale

    # Neuroticism Scaler
    if type == "Neuroticism":
        # N_scaler=MinMaxScaler(feature_range=(-3.0, 2.0))
        N_score_data = N_scaler.fit_transform(score)
        N_score_scale = N_score_data[result]
        return N_score_scale


O_scaler = MinMaxScaler(feature_range=(-3.2, 2.3))
C_scaler = MinMaxScaler(feature_range=(-2.5, 2.4))
E_scaler = MinMaxScaler(feature_range=(-3.6, 2.5))
A_scaler = MinMaxScaler(feature_range=(-3.0, 2.1))
N_scaler = MinMaxScaler(feature_range=(-3.0, 2.0))


def Avg_Inverse_Result(score_scale):
    E_result = score_scale[0]
    A_result = score_scale[1]
    C_result = score_scale[2]
    N_result = score_scale[3]
    O_result = score_scale[4]

    # Inverse scale
    E_final = E_scaler.inverse_transform([[E_result], [0]])
    A_final = A_scaler.inverse_transform([[A_result], [0]])
    C_final = C_scaler.inverse_transform([[C_result], [0]])
    N_final = N_scaler.inverse_transform([[N_result], [0]])
    O_final = O_scaler.inverse_transform([[O_result], [0]])

    OCEAN_final = []
    OCEAN_final.append(E_final[0])
    OCEAN_final.append(A_final[0])
    OCEAN_final.append(C_final[0])
    OCEAN_final.append(N_final[0])
    OCEAN_final.append(O_final[0])

    return OCEAN_final


def create_docx_if_not_exist(file_path):
    if not os.path.exists(file_path):
        doc = Document()
        doc.save(file_path)
        print(f"Created a new Word document: {file_path}")
    else:
        print(f"The Word document already exists: {file_path}")


def convert_docx_to_pdf(docx_file, pdf_file):
    try:
        pyuno_path = "/usr/bin/libreoffice/program/pyuno.so"
        python_binary = "/root/iscv/machine/myconda/bin/python3"
        # command = [
        #     "python3",
        #     "-m",
        #     "unoconv",
        #     "--pyuno",
        #     pyuno_path,
        #     "-f",
        #     "pdf",
        #     "-o",
        #     pdf_file,
        #     docx_file,
        # ]
        # subprocess.run(command)
        subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                pdf_file,
                docx_file,
            ]
        )
        # convert(docx_file, pdf_file)
        print("Conversion successful!")
    except Exception as e:
        print(f"An error occurred: {str(e)}")


async def read_qa(path):
    with open(path, "r") as file:
        lines = file.readlines()

    result = []
    item_count = 1

    # Tạo danh sách từ điển với item tăng dần và value lấy từ file .txt
    for line in lines:
        value = (
            line.strip()
        )  # Lấy giá trị từ file .txt, xoá ký tự trống ở đầu và cuối dòng
        if value:  # Kiểm tra nếu giá trị không rỗng
            data = {"item": item_count, "value": value}
            result.append(data)
            item_count += 1
    # In danh sách từ điển
    return result


async def handle_report(data):
    session_id = data["sessionId"]
    current_path = os.path.dirname(__file__) + "/"
    folder_path = f"./public/interview/{session_id}/"
    qa_path = folder_path + "qa.txt"
    avgscores = average_big_five(data["bigfive"])
    task1 = asyncio.create_task(drawGraph(avgscores, folder_path))
    task2 = asyncio.create_task(BigFiveComment(avgscores, current_path))
    task3 = asyncio.create_task(read_qa(qa_path))
    results = await asyncio.gather(task1, task2, task3)
    comments = results[1]
    qa = results[2]
    input = {
        "employeeId": data["employeeId"],
        "employeeName": data["employeeName"],
        "o": data["bigfive"]["o"],
        "c": data["bigfive"]["c"],
        "e": data["bigfive"]["e"],
        "a": data["bigfive"]["a"],
        "n": data["bigfive"]["n"],
        "oc": comments["oc"],
        "cc": comments["cc"],
        "ec": comments["ec"],
        "ac": comments["ac"],
        "nc": comments["nc"],
        "question": qa,
    }
    fill_word(input, current_path, folder_path)
    convert_docx_to_pdf(folder_path + "report.docx", folder_path)


def radar_factory(num_vars, frame="circle"):
    theta = np.linspace(0, 2 * np.pi, num_vars, endpoint=False)

    class RadarAxes(PolarAxes):
        name = "radar"

        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            # rotate plot such that the first axis is at the top
            self.set_theta_zero_location("N")

        def fill(self, *args, closed=True, **kwargs):
            """Override fill so that line is closed by default"""
            return super().fill(closed=closed, *args, **kwargs)

        def plot(self, *args, **kwargs):
            """Override plot so that line is closed by default"""
            lines = super().plot(*args, **kwargs)
            for line in lines:
                self._close_line(line)

        def _close_line(self, line):
            x, y = line.get_data()
            # FIXME: markers at x[0], y[0] get doubled-up
            if x[0] != x[-1]:
                x = np.concatenate((x, [x[0]]))
                y = np.concatenate((y, [y[0]]))
                line.set_data(x, y)

        def set_varlabels(self, labels):
            self.set_thetagrids(np.degrees(theta), labels)

        def _gen_axes_patch(self):
            # The Axes patch must be centered at (0.5, 0.5) and of radius 0.5
            # in axes coordinates.
            if frame == "circle":
                return Circle((0.5, 0.5), 0.5)
            elif frame == "polygon":
                return RegularPolygon((0.5, 0.5), num_vars, radius=0.5, edgecolor="k")
            else:
                raise ValueError("unknown value for 'frame': %s" % frame)

        def draw(self, renderer):
            """Draw. If frame is polygon, make gridlines polygon-shaped"""
            if frame == "polygon":
                gridlines = self.yaxis.get_gridlines()
                for gl in gridlines:
                    gl.get_path()._interpolation_steps = num_vars
            super().draw(renderer)

        def _gen_axes_spines(self):
            if frame == "circle":
                return super()._gen_axes_spines()
            elif frame == "polygon":
                # spine_type must be 'left'/'right'/'top'/'bottom'/'circle'.
                spine = Spine(
                    axes=self,
                    spine_type="circle",
                    path=Path.unit_regular_polygon(num_vars),
                )
                # unit_regular_polygon gives a polygon of radius 1 centered at
                # (0, 0) but we want a polygon of radius 0.5 centered at (0.5,
                # 0.5) in axes coordinates.
                spine.set_transform(
                    Affine2D().scale(0.5).translate(0.5, 0.5) + self.transAxes
                )
                return {"polar": spine}
            else:
                raise ValueError("unknown value for 'frame': %s" % frame)

    register_projection(RadarAxes)
    return theta


def average_big_five(bigfive):
    data = [
        [
            "Extroversion",
            "Agreeableness",
            "Conscientiousness",
            "Neuroticism",
            "Openness",
        ],
        (
            "The Big Five Personality Score",
            [
                [
                    bigfive["e"],
                    bigfive["a"],
                    bigfive["c"],
                    bigfive["n"],
                    bigfive["o"],
                ]
            ],
        ),
    ]
    return data


async def drawGraph(data, path):
    N = len(data[0])
    theta = radar_factory(N, frame="polygon")  # polygon  !!!

    spoke_labels = data.pop(0)
    title, case_data = data[0]
    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(projection="radar"))
    fig.subplots_adjust(top=0.85, bottom=0.05)
    ax.set_rgrids([0, 10, 20, 30, 40])
    ax.set_title(title, position=(0.5, 1.1), ha="center")
    ax.set_ylim(0, 40)
    for d in case_data:
        line = ax.plot(theta, d)
        ax.fill(theta, d, alpha=0.25)
    ax.set_varlabels(spoke_labels)

    plt.savefig(path + "chart.png")
    plt.close()
    # plt.show()


# =================================================================
async def BigFiveComment(result, current_path):
    # Lưu số điểm từng tính cách vào dataframe
    df_result = pd.DataFrame(
        (
            [
                ["Extroversion", result[0][1][0][0]],
                ["Agreeableness", result[0][1][0][1]],
                ["Conscientiousness", result[0][1][0][2]],
                ["Neuroticism", result[0][1][0][3]],
                ["Openness to Experience", result[0][1][0][4]],
            ]
        ),
        columns=["Type", "Result"],
    )

    # Lưu dữ liệu đánh giá tính cách vào dataframe
    df_comment = pd.read_csv(current_path + "comments.csv")

    # Merge 2 dataframe
    df_merge = pd.merge(df_comment, df_result, on="Type", how="outer")
    # Lọc ra những dòng dữ liệu thỏa điều kiện
    df_final = df_merge[
        (df_merge["Result"] >= df_merge["StartScore"])
        & (df_merge["Result"] <= df_merge["EndScore"])
    ].reset_index()
    # Hiển thị đánh giá

    # for i in range(5):
    #     comments.append(
    #         df_final["Type"][i]
    #         + " Comment: "
    #         + to_str(df_final["Comment"][i]).replace("'", "")
    #     )
    result = {}
    result["ec"] = df_final["Comment"][0]
    result["ac"] = df_final["Comment"][1]
    result["cc"] = df_final["Comment"][2]
    result["nc"] = df_final["Comment"][3]
    result["oc"] = df_final["Comment"][4]
    return result


def fill_word(dictionary, current_path, folder_path):
    document = Document(current_path + "template.docx")
    paragraph_index = 3
    if paragraph_index < len(document.paragraphs):
        # Truy cập vào đoạn văn cần xem
        paragraph = document.paragraphs[paragraph_index]
        text = paragraph.text

        paragraph.text = text[:9] + "{}".format(dictionary["employeeId"]) + text[9:]
        # In ra nội dung của đoạn văn
        text = paragraph.text

        for run in paragraph.runs:
            run.font.name = "Times New Roman"
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.bold = True
            words = text.split()
            for word in words:
                # Kiểm tra nếu từ là "ID001"
                if word == "-":
                    text = paragraph.add_run(
                        "Người thực hiện phỏng vấn: " + dictionary["employeeName"] + " "
                    )
                    # Đặt chữ in đậm
                    text.font.italic = True
                    text.font.name = "Times New Roman"

            # nó in nghiên cái ID luôn rồi a

    # Load dữ liệu vào bảng

    table_index = 1
    table = document.tables[table_index]
    for i in range(len(table.rows) - 1):
        row_index = i + 1
        column_index1 = 2
        cell = table.cell(row_index, column_index1)
        # Thêm dữ liệu vào ô
        paragraph1 = cell.add_paragraph()
        # Thêm dữ liệu từ dictionary
        value = str(dictionary["question"][i]["value"])
        run = paragraph1.add_run(value)
        if cell.text.startswith(""):
            cell.text = cell.text.lstrip()

    for i in range(len(table.rows) - 1):
        row_index = i + 1
        column_index1 = 6
        cell = table.cell(row_index, column_index1)
        k = i + 25
        # Thêm dữ liệu vào ô
        paragraph1 = cell.add_paragraph()
        # Thêm dữ liệu từ dictionary
        value1 = str(dictionary["question"][k]["value"])
        run = paragraph1.add_run((value1))
        if cell.text.startswith(""):
            cell.text = cell.text.lstrip()
    # Định dạng dữ liệu
    column_index = 2
    for row in table.rows:
        cell = row.cells[column_index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

        # In đậm dữ liệu và căn giữa trong ô
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                paragraph.alignment = WD_ALIGN_VERTICAL.CENTER
                run.font.name = "Times New Roman"

    column_index = 6
    for row in table.rows:
        cell = row.cells[column_index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

        # In đậm dữ liệu và căn giữa trong ô
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                paragraph.alignment = WD_ALIGN_VERTICAL.CENTER
                run.font.name = "Times New Roman"
    # Thêm hình ảnh
    from docx.shared import Inches
    from docx.shared import Cm

    image_path = folder_path + "chart.png"
    tables = document.tables
    table_index = 2
    row_index = 0
    column_index = 0
    table = tables[table_index]
    cell = table.cell(row_index, column_index)
    paragraph2 = cell.paragraphs[0]
    run = paragraph2.add_run()
    run.add_picture(image_path, width=Cm(5), height=Cm(5))
    paragraph2.alignment = 1

    # Thêm chỉ số

    table_index = 2
    column_index = 2
    # Openness
    row_index = 1
    table = tables[table_index]
    cell = table.cell(row_index, column_index)

    # Thêm chỉ số từ dictionary
    o = str(dictionary["o"])
    o += "/40"
    paragraph3 = cell.add_paragraph()
    run = paragraph3.add_run(o)
    if cell.text.startswith(""):
        cell.text = cell.text.lstrip()
    # Conscientiousness
    row_index = 2
    table = tables[table_index]
    cell = table.cell(row_index, column_index)
    # Thêm chỉ số từ dictionary
    c = str(dictionary["c"])
    c += "/40"
    paragraph3 = cell.add_paragraph()
    run = paragraph3.add_run(c)
    if cell.text.startswith(""):
        cell.text = cell.text.lstrip()
    # Extroversion
    row_index = 3
    table = tables[table_index]
    cell = table.cell(row_index, column_index)
    # Thêm chỉ số từ dictionary
    e = str(dictionary["e"])
    e += "/40"
    paragraph3 = cell.add_paragraph()
    run = paragraph3.add_run(e)
    if cell.text.startswith(""):
        cell.text = cell.text.lstrip()
    # Agreeableness
    row_index = 4
    table = tables[table_index]
    cell = table.cell(row_index, column_index)
    # Thêm chỉ số từ dictionary
    a = str(dictionary["a"])
    a += "/40"
    paragraph3 = cell.add_paragraph()
    run = paragraph3.add_run(a)
    if cell.text.startswith(""):
        cell.text = cell.text.lstrip()
    # Neuroticism
    row_index = 5
    table = tables[table_index]
    cell = table.cell(row_index, column_index)

    # Thêm chỉ số từ dictionary
    n = str(dictionary["n"])
    n += "/40"
    paragraph3 = cell.add_paragraph()
    run = paragraph3.add_run(n)
    if cell.text.startswith(""):
        cell.text = cell.text.lstrip()

    column_index = 2
    for row in table.rows:
        cell = row.cells[column_index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

        # In đậm dữ liệu và căn giữa trong ô
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                paragraph.alignment = WD_ALIGN_VERTICAL.CENTER
                run.font.name = "Times New Roman"

    # Thêm kết luận

    # Lấy đoạn Extroversion trong docx
    paragraph_index = 8
    paragraph = document.paragraphs[paragraph_index]
    # Lấy Extroversion Comment trong txt

    # Thêm kết luận từ dictionary
    ec = dictionary["ec"]
    text = paragraph.add_run(ec)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"

    # Lấy đoạn Conscientiousness trong docx
    paragraph_index = 9
    paragraph = document.paragraphs[paragraph_index]
    # Lấy Conscientiousness Comment trong txt

    # Thêm kết luận từ dictionary
    cc = dictionary["cc"]
    text = paragraph.add_run(cc)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"

    # Lấy đoạn Openness to Experience trong docx
    paragraph_index = 10
    paragraph = document.paragraphs[paragraph_index]
    # Lấy Openness to Experience Comment trong txt

    # Thêm kết luận từ dictionary
    oc = dictionary["oc"]
    text = paragraph.add_run(oc)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"

    # Lấy đoạn Agreeableness trong docx
    paragraph_index = 11
    paragraph = document.paragraphs[paragraph_index]
    # Lấy Openness to Agreeableness Comment trong txt

    # Thêm kết luận từ dictionary
    ac = dictionary["ac"]
    text = paragraph.add_run(ac)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"

    # Lấy đoạn Neuroticism trong docx
    paragraph_index = 12
    paragraph = document.paragraphs[paragraph_index]
    # Lấy Openness to Neuroticism Comment trong txt

    # Thêm kết luận từ dictionary
    nc = dictionary["ec"]
    text = paragraph.add_run(nc)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
    # Thêm Ngày Tháng Năm Hiện tại
    paragraph_index = 13
    paragraph = document.paragraphs[paragraph_index]
    # Ngày tháng năm hiện tại
    day = datetime.now().day
    month = datetime.now().month
    year = datetime.now().year

    text = paragraph.text
    # Tạo chuỗi ngày tháng năm
    ngay_thang_nam = f"{day} tháng {month} năm {year}"
    paragraph.text += ngay_thang_nam
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
    for run in paragraph.runs:
        run.font.size = Pt(11)

    document.save(folder_path + "report.docx")
